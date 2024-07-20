import docx

doc = docx.Document("test_word.docx")

print("段落數量:"+ str(len(doc.paragraphs)))

TEXT = doc.paragraphs

doc.paragraphs[0].add_run("熊先生\n")
doc.add_paragraph("日期20240720")
doc.save("new_file.docx")

#print(doc.paragraphs[0].text)

#for i in range(len(doc.paragraphs)):

#    print(TEXT[i].text)
"""
for paragraph in doc.paragraphs:
    print(paragraph.text)


tune = 0

while tune < len(doc.paragraphs):

    print(doc.paragraphs[tune].text)
    tune += 1

"""

